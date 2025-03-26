from flask import Flask, render_template, request, send_file
import os
import csv
from docx import Document
from crossref.restful import Works

app = Flask(__name__)

def write_reference(doi, document):
    works = Works()
    paper = works.doi(doi)
    if paper is None:
        return
    title = paper['title'][0]
    journal = paper['container-title'][0]
    year = paper['created']['date-parts'][0][0]
    authors = ', '.join(f"{author['family']} {author['given'][0]}" for author in paper['author'])

    p = document.add_paragraph(style='List Bullet')
    p = document.add_paragraph()
    p.add_run(f"{title}. ").bold = True
    p.add_run(f"{authors} {journal}, {year}. doi: {doi}")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'csv_file' not in request.files:
        return 'No file part', 400

    file = request.files['csv_file']
    if file.filename == '':
        return 'No selected file', 400

    # ƒtƒ@ƒCƒ‹‚Ì•Û‘¶
    filename = os.path.join('uploads', file.filename)
    file.save(filename)

    # CSV‚ğˆ—‚µ‚ÄWord•¶‘‚ğ¶¬
    output_docx = 'output.docx'
    document = Document()
    with open(filename, encoding='utf-8') as f:
        reader = csv.reader(f)
        for row in reader:
            write_reference(row[0], document)
    document.save(output_docx)

    return send_file(output_docx, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
